from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture('me.png',
                     width=Inches(2.0),
                     height=Inches(2.0))

# name, phone number, and email details
name = input('What is your name? ')
speak('Hello' + name + 'How are you today? ')

speak('What is your phone number? ')
phoneNumber = input('What is your phone number? ')

speak('What is your email? ')
email = input('What is your email? ')

document.add_paragraph(name + ' | ' + phoneNumber + ' | ' + email)

# about me
document.add_heading('About me')
document.add_paragraph(input("Tell us about yourself? "))

# Work experience
document.add_heading('Education')
p = document.add_paragraph()

university = input('Enter university name ')
from_date = input('From date ')
to_date = input('To date ')

# add text to existing paragraph
p.add_run(university + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your education at ' + university + ' ')

p.add_run(experience_details)

# more experiences
while True:
    hasmoreExperiences = input(
        'Do you have more Experiences? Yes or No? ')

    if hasmoreExperiences.lower() == 'yes':

        p = document.add_paragraph

        company = input('Enter company ')
        from_date = input('From date ')
        to_date = input('To date ')

        # add text to existing parag
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ' ')
        p.add_run(experience_details)

    else:
        break

# skills
document.add_heading('Skills')
skill = input('Enter skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    hasmoreSkills = input('Has more skills? Yes or No? ')
    if hasmoreSkills.lower() == 'yes':
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
speak("CV generated, Thank you for using this program")
p.text = "CV generated"

document.save('cv.docx')
